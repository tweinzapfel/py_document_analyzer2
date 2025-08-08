import streamlit as st
import pandas as pd
from datetime import datetime
import re
import io
import base64
from typing import List, Dict, Tuple
import PyPDF2
import docx
from openpyxl import Workbook
from openpyxl.styles import Font, Fill, PatternFill, Alignment
from openpyxl.utils.dataframe import dataframe_to_rows

# Page configuration
st.set_page_config(
    page_title="RFP Document Analyzer",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        color: #1e3a8a;
        border-bottom: 3px solid #3b82f6;
        padding-bottom: 10px;
        margin-bottom: 30px;
    }
    .risk-high {
        background-color: #fee2e2;
        padding: 10px;
        border-left: 4px solid #dc2626;
        margin: 10px 0;
    }
    .risk-medium {
        background-color: #fef3c7;
        padding: 10px;
        border-left: 4px solid #f59e0b;
        margin: 10px 0;
    }
    .risk-low {
        background-color: #dbeafe;
        padding: 10px;
        border-left: 4px solid #3b82f6;
        margin: 10px 0;
    }
    .instruction-item {
        background-color: #f0f9ff;
        padding: 12px;
        margin: 8px 0;
        border-radius: 6px;
        border-left: 3px solid #0284c7;
    }
    .deadline-alert {
        background-color: #fef2f2;
        padding: 8px;
        border-radius: 4px;
        font-weight: bold;
        color: #991b1b;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'custom_terms' not in st.session_state:
    st.session_state.custom_terms = []

# Risk keywords and patterns
RISK_PATTERNS = {
    'high': {
        'unlimited_liability': r'unlimited\s+liability|uncapped\s+liability',
        'indemnification': r'indemnif|hold\s+harmless',
        'ip_transfer': r'work\s+for\s+hire|assigns?\s+all\s+rights?|transfer\s+of\s+ownership',
        'liquidated_damages': r'liquidated\s+damages|penalty\s+clause',
        'termination_for_convenience': r'termination\s+for\s+convenience|terminate\s+without\s+cause',
        'payment_terms': r'net\s+\d{2,3}|payment\s+terms?\s+\d{2,}|paid\s+after\s+\d{2,}',
    },
    'medium': {
        'warranty': r'warrant[yi]|guarantee',
        'insurance': r'insurance\s+requirements?|liability\s+insurance|errors?\s+and\s+omissions?',
        'audit_rights': r'audit\s+rights?|right\s+to\s+audit|inspection\s+rights?',
        'confidentiality': r'confidential|non-?disclosure|proprietary\s+information',
        'dispute_resolution': r'arbitration|dispute\s+resolution|governing\s+law',
    },
    'low': {
        'delivery_terms': r'delivery\s+date|milestone|deliverable',
        'acceptance_criteria': r'acceptance\s+criteria|acceptance\s+testing',
        'change_orders': r'change\s+order|modification|amendment',
    }
}

# Federal-specific patterns
FEDERAL_PATTERNS = {
    'far_clauses': r'FAR\s+\d+\.\d+|Federal\s+Acquisition\s+Regulation',
    'dfars_clauses': r'DFARS\s+\d+\.\d+|Defense\s+Federal\s+Acquisition',
    'small_business': r'small\s+business|8\(a\)|HUBZone|SDVOSB|WOSB',
    'security_clearance': r'security\s+clearance|classified|secret|confidential',
    'buy_american': r'Buy\s+American|domestic\s+end\s+product|TAA\s+compliant',
}

# Instruction patterns
INSTRUCTION_PATTERNS = {
    'deadline': r'due\s+date|deadline|submit\s+by|submission\s+date|closing\s+date|must\s+be\s+received',
    'format': r'page\s+limit|font\s+size|margin|spacing|format\s+requirement',
    'submission_method': r'submit\s+to|email\s+to|upload\s+to|deliver\s+to|submission\s+portal',
    'required_docs': r'required\s+document|must\s+include|shall\s+provide|attachment|exhibit',
    'evaluation': r'evaluation\s+criteria|scoring|weight|technical\s+approach|past\s+performance',
}

def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def find_matches(text: str, patterns: Dict[str, str]) -> Dict[str, List[str]]:
    """Find pattern matches in text"""
    matches = {}
    for key, pattern in patterns.items():
        found = re.findall(pattern, text, re.IGNORECASE)
        if found:
            # Get context around matches
            contexts = []
            for match in set(found[:5]):  # Limit to 5 examples
                pattern_with_context = r'.{0,100}' + re.escape(match) + r'.{0,100}'
                context_matches = re.findall(pattern_with_context, text, re.IGNORECASE)
                if context_matches:
                    contexts.append(context_matches[0].strip())
            matches[key] = contexts
    return matches

def extract_deadlines(text: str) -> List[str]:
    """Extract potential deadline information"""
    deadlines = []
    # Look for date patterns
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{2,4}\b',
    ]
    
    for pattern in date_patterns:
        dates = re.findall(pattern, text, re.IGNORECASE)
        for date in dates:
            # Get context around the date
            context_pattern = r'.{0,50}' + re.escape(date) + r'.{0,50}'
            contexts = re.findall(context_pattern, text, re.IGNORECASE)
            for context in contexts:
                if any(word in context.lower() for word in ['due', 'deadline', 'submit', 'close', 'by']):
                    deadlines.append(context.strip())
    
    return list(set(deadlines))[:10]  # Return unique, limited to 10

def analyze_document(text: str, custom_terms: List[str]) -> Dict:
    """Perform comprehensive document analysis"""
    results = {
        'instructions': {},
        'risks': {'high': {}, 'medium': {}, 'low': {}},
        'federal_items': {},
        'custom_searches': {},
        'deadlines': [],
        'statistics': {}
    }
    
    # Extract instructions
    results['instructions'] = find_matches(text, INSTRUCTION_PATTERNS)
    
    # Extract deadlines
    results['deadlines'] = extract_deadlines(text)
    
    # Identify risks
    for risk_level, patterns in RISK_PATTERNS.items():
        results['risks'][risk_level] = find_matches(text, patterns)
    
    # Check for federal-specific items
    results['federal_items'] = find_matches(text, FEDERAL_PATTERNS)
    
    # Search for custom terms
    if custom_terms:
        custom_patterns = {term: re.escape(term) for term in custom_terms}
        results['custom_searches'] = find_matches(text, custom_patterns)
    
    # Calculate statistics
    results['statistics'] = {
        'total_pages': text.count('\n') // 50,  # Rough estimate
        'word_count': len(text.split()),
        'risk_count': sum(len(risks) for level in results['risks'].values() for risks in level.values()),
        'instruction_count': sum(len(instr) for instr in results['instructions'].values()),
    }
    
    return results

def create_excel_report(results: Dict) -> bytes:
    """Create Excel report from analysis results"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Summary sheet
        summary_data = {
            'Metric': ['Total Word Count', 'Estimated Pages', 'Total Risks Identified', 
                      'Total Instructions Found', 'Deadlines Found'],
            'Value': [
                results['statistics']['word_count'],
                results['statistics']['total_pages'],
                results['statistics']['risk_count'],
                results['statistics']['instruction_count'],
                len(results['deadlines'])
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        # Instructions sheet
        instruction_rows = []
        for category, items in results['instructions'].items():
            for item in items:
                instruction_rows.append({
                    'Category': category.replace('_', ' ').title(),
                    'Instruction': item
                })
        if instruction_rows:
            instructions_df = pd.DataFrame(instruction_rows)
            instructions_df.to_excel(writer, sheet_name='Instructions', index=False)
        
        # Risks sheet
        risk_rows = []
        for level in ['high', 'medium', 'low']:
            for risk_type, items in results['risks'][level].items():
                for item in items:
                    risk_rows.append({
                        'Risk Level': level.upper(),
                        'Risk Type': risk_type.replace('_', ' ').title(),
                        'Context': item
                    })
        if risk_rows:
            risks_df = pd.DataFrame(risk_rows)
            risks_df.to_excel(writer, sheet_name='Risks', index=False)
        
        # Deadlines sheet
        if results['deadlines']:
            deadlines_df = pd.DataFrame({'Deadline Information': results['deadlines']})
            deadlines_df.to_excel(writer, sheet_name='Deadlines', index=False)
        
        # Federal items sheet (if applicable)
        if results['federal_items']:
            federal_rows = []
            for category, items in results['federal_items'].items():
                for item in items:
                    federal_rows.append({
                        'Category': category.replace('_', ' ').title(),
                        'Reference': item
                    })
            if federal_rows:
                federal_df = pd.DataFrame(federal_rows)
                federal_df.to_excel(writer, sheet_name='Federal Requirements', index=False)
        
        # Custom searches sheet
        if results['custom_searches']:
            custom_rows = []
            for term, items in results['custom_searches'].items():
                for item in items:
                    custom_rows.append({
                        'Search Term': term,
                        'Context': item
                    })
            if custom_rows:
                custom_df = pd.DataFrame(custom_rows)
                custom_df.to_excel(writer, sheet_name='Custom Searches', index=False)
        
        # Format the workbook
        workbook = writer.book
        for worksheet in workbook.worksheets:
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    return output.getvalue()

def display_html_results(results: Dict):
    """Display results in HTML format"""
    st.markdown("## 📋 Analysis Results")
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Word Count", f"{results['statistics']['word_count']:,}")
    with col2:
        st.metric("Risks Found", results['statistics']['risk_count'])
    with col3:
        st.metric("Instructions", results['statistics']['instruction_count'])
    with col4:
        st.metric("Deadlines", len(results['deadlines']))
    
    # Create tabs for different sections
    tabs = st.tabs(["📝 Instructions", "⚠️ Risks", "📅 Deadlines", "🏛️ Federal Items", "🔍 Custom Searches"])
    
    # Instructions tab
    with tabs[0]:
        st.markdown("### Submission Instructions")
        if results['instructions']:
            for category, items in results['instructions'].items():
                if items:
                    st.markdown(f"**{category.replace('_', ' ').title()}**")
                    for item in items:
                        st.markdown(f'<div class="instruction-item">{item}</div>', unsafe_allow_html=True)
        else:
            st.info("No specific instructions found")
    
    # Risks tab
    with tabs[1]:
        st.markdown("### Risk Assessment")
        
        # High risks
        if results['risks']['high']:
            st.markdown("#### 🔴 High Risk Items")
            for risk_type, items in results['risks']['high'].items():
                if items:
                    st.markdown(f"**{risk_type.replace('_', ' ').title()}**")
                    for item in items:
                        st.markdown(f'<div class="risk-high">{item}</div>', unsafe_allow_html=True)
        
        # Medium risks
        if results['risks']['medium']:
            st.markdown("#### 🟡 Medium Risk Items")
            for risk_type, items in results['risks']['medium'].items():
                if items:
                    st.markdown(f"**{risk_type.replace('_', ' ').title()}**")
                    for item in items:
                        st.markdown(f'<div class="risk-medium">{item}</div>', unsafe_allow_html=True)
        
        # Low risks
        if results['risks']['low']:
            st.markdown("#### 🔵 Low Risk Items")
            for risk_type, items in results['risks']['low'].items():
                if items:
                    st.markdown(f"**{risk_type.replace('_', ' ').title()}**")
                    for item in items:
                        st.markdown(f'<div class="risk-low">{item}</div>', unsafe_allow_html=True)
        
        if not any(results['risks'].values()):
            st.info("No specific risks identified")
    
    # Deadlines tab
    with tabs[2]:
        st.markdown("### Important Deadlines")
        if results['deadlines']:
            for deadline in results['deadlines']:
                st.markdown(f'<div class="deadline-alert">📅 {deadline}</div>', unsafe_allow_html=True)
        else:
            st.info("No specific deadlines found")
    
    # Federal items tab
    with tabs[3]:
        st.markdown("### Federal Contract Requirements")
        if results['federal_items']:
            for category, items in results['federal_items'].items():
                if items:
                    st.markdown(f"**{category.replace('_', ' ').title()}**")
                    for item in items:
                        st.markdown(f"- {item}")
        else:
            st.info("No federal-specific requirements identified")
    
    # Custom searches tab
    with tabs[4]:
        st.markdown("### Custom Term Search Results")
        if results['custom_searches']:
            for term, items in results['custom_searches'].items():
                if items:
                    st.markdown(f"**Search term: '{term}'**")
                    for item in items:
                        st.markdown(f"- {item}")
        else:
            st.info("No custom search terms provided or no matches found")

# Main app
def main():
    st.markdown('<h1 class="main-header">📄 RFP Document Analyzer</h1>', unsafe_allow_html=True)
    st.markdown("Upload RFP documents to extract submission instructions, identify contractual risks, and generate comprehensive reports.")
    
    # Sidebar for configuration
    with st.sidebar:
        st.markdown("### ⚙️ Configuration")
        
        # Custom search terms
        st.markdown("#### Custom Search Terms")
        new_term = st.text_input("Add custom term to search for:")
        if st.button("Add Term") and new_term:
            st.session_state.custom_terms.append(new_term)
        
        if st.session_state.custom_terms:
            st.markdown("**Current search terms:**")
            for i, term in enumerate(st.session_state.custom_terms):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.text(term)
                with col2:
                    if st.button("❌", key=f"remove_{i}"):
                        st.session_state.custom_terms.pop(i)
                        st.rerun()
        
        st.markdown("---")
        st.markdown("### 📊 Risk Categories")
        st.markdown("""
        **🔴 High Risk**
        - Unlimited liability
        - IP transfer
        - Liquidated damages
        
        **🟡 Medium Risk**
        - Warranties
        - Insurance requirements
        - Audit rights
        
        **🔵 Low Risk**
        - Delivery terms
        - Acceptance criteria
        - Change orders
        """)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 📤 Upload Documents")
        uploaded_files = st.file_uploader(
            "Choose RFP documents (PDF or DOCX)",
            type=['pdf', 'docx'],
            accept_multiple_files=True,
            help="You can upload multiple RFP documents for analysis"
        )
    
    with col2:
        st.markdown("### 🎯 Analysis Options")
        analyze_federal = st.checkbox("Include Federal Contract Analysis", value=True)
        include_statistics = st.checkbox("Include Document Statistics", value=True)
    
    if uploaded_files:
        st.markdown("---")
        
        # Show uploaded files
        st.markdown("### 📁 Uploaded Files")
        file_info = []
        for file in uploaded_files:
            file_info.append({
                "File Name": file.name,
                "File Type": file.type,
                "File Size": f"{file.size / 1024:.2f} KB"
            })
        st.dataframe(pd.DataFrame(file_info), use_container_width=True)
        
        # Analyze button
        if st.button("🔍 Analyze Documents", type="primary", use_container_width=True):
            with st.spinner("Analyzing documents... This may take a moment."):
                all_text = ""
                
                # Extract text from all uploaded files
                progress = st.progress(0)
                for i, file in enumerate(uploaded_files):
                    progress.progress((i + 1) / len(uploaded_files))
                    
                    if file.type == "application/pdf":
                        text = extract_text_from_pdf(file)
                    else:  # DOCX
                        text = extract_text_from_docx(file)
                    
                    all_text += f"\n\n--- Document: {file.name} ---\n\n" + text
                
                # Perform analysis
                if all_text.strip():
                    results = analyze_document(all_text, st.session_state.custom_terms)
                    st.session_state.analysis_results = results
                    st.success("✅ Analysis complete!")
                else:
                    st.error("❌ Could not extract text from the uploaded documents.")
    
    # Display results if available
    if st.session_state.analysis_results:
        st.markdown("---")
        
        # Display HTML results
        display_html_results(st.session_state.analysis_results)
        
        # Download button for Excel report
        st.markdown("---")
        st.markdown("### 💾 Export Report")
        
        excel_data = create_excel_report(st.session_state.analysis_results)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        st.download_button(
            label="📥 Download Excel Report",
            data=excel_data,
            file_name=f"RFP_Analysis_Report_{timestamp}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #6b7280; font-size: 0.9em;'>
    RFP Document Analyzer | Supports PDF and DOCX formats | Identifies risks and extracts submission requirements
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
